#!/usr/bin/env python3
"""
DevSecOps Handbook DOCX Generator
สร้างไฟล์ DevSecOps-Handbook.docx อัตโนมัติจากเนื้อหา Markdown
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

# ============== CONFIGURATION ==============
TITLE = "DevSecOps Handbook"
SUBTITLE = "คู่มือมาตรฐานความปลอดภัย, การพัฒนา และการดำเนินงานอัตโนมัติ"
VERSION = "1.0"
AUTHOR = "DevSecOps Team"
DATE = "2026"
OUTPUT_FILE = "DevSecOps-Handbook.docx"
# ============================================

def set_heading_style(paragraph, level):
    """ตั้งค่าสไตล์หัวข้อ"""
    run = paragraph.runs[0]
    run.bold = True
    run.font.size = Pt(24 - (level * 2))
    run.font.color.rgb = RGBColor(0, 51, 102)
    paragraph.space_after = Pt(12)

def add_cover_page(doc):
    """เพิ่มหน้าปก"""
    doc.add_paragraph("\n" * 5)
    title = doc.add_heading(TITLE, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_heading_style(title, 0)
    
    subtitle = doc.add_paragraph(SUBTITLE)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].italic = True
    subtitle.space_after = Pt(40)
    
    info = doc.add_paragraph(f"Version {VERSION} — {DATE}")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.size = Pt(12)
    
    author = doc.add_paragraph(f"Author: {AUTHOR}")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.runs[0].font.size = Pt(12)
    
    doc.add_page_break()

def add_toc_placeholder(doc):
    """เพิ่มสารบัญ (Word จะอัพเดทเอง)"""
    doc.add_heading("สารบัญ", level=1)
    toc = doc.add_paragraph()
    toc_run = toc.add_run("«คลิกขวา → Update Field เพื่ออัพเดทสารบัญ»")
    toc_run.italic = True
    toc_run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()

def add_section(doc, title, content_en="", content_th="", level=1):
    """เพิ่มส่วนเนื้อหา"""
    heading = doc.add_heading(title, level=level)
    set_heading_style(heading, level)
    
    if content_en:
        en_para = doc.add_paragraph()
        en_label = en_para.add_run("EN: ")
        en_label.bold = True
        en_label.font.color.rgb = RGBColor(50, 50, 150)
        en_para.add_run(content_en)
        en_para.space_after = Pt(6)
    
    if content_th:
        th_para = doc.add_paragraph()
        th_label = th_para.add_run("TH: ")
        th_label.bold = True
        th_label.font.color.rgb = RGBColor(150, 50, 50)
        th_para.add_run(content_th)
        th_para.space_after = Pt(12)

def add_index(doc):
    """เพิ่มดัชนีคำสำคัญ"""
    doc.add_page_break()
    doc.add_heading("ดัชนีคำสำคัญ (Index)", level=1)
    
    index_data = [
        ("Access Control", "6"), ("Advisory", "1"), ("AI Agents", "31"),
        ("API", "11, 29"), ("Audit Trail", "23"), ("Authentication", "6, 9"),
        ("Automation", "24"), ("Backup", "30"), ("CI/CD", "10, 26"),
        ("CodeQL", "27"), ("Compliance", "22"), ("CVSS", "28"),
        ("Dashboard", "29"), ("DataService", "29"), ("Database", "30"),
        ("Debugging", "26"), ("GDPR", "22"), ("GitLab CI", "10"),
        ("Go", "14"), ("Incident Response", "28"), ("ISO 27001", "22"),
        ("Java", "19"), ("JavaScript", "15"), ("Jira", "24"),
        ("Knowledge Base", "28"), ("License", "3"), ("Obsidian", "21, 28"),
        ("Ollama", "21, 31"), ("Policies", "6"), ("Privacy", "4"),
        ("Python", "13"), ("REST API", "11"), ("Rust", "18"),
        ("SAST/DAST", "27"), ("Scripts", "2"), ("Security Scan", "27"),
        ("Sigma", "17"), ("Slack", "24"), ("Snyk", "27"),
        ("SOC2", "22"), ("Terms of Service", "5"), ("Threat Model", "28"),
        ("Trivy", "27"), ("Webhook", "9"), ("WebUI", "8"),
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "คำสำคัญ"
    hdr_cells[1].text = "ส่วน"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    for keyword, section in index_data:
        row_cells = table.add_row().cells
        row_cells[0].text = keyword
        row_cells[1].text = section

def main():
    print("🚀 กำลังสร้าง DevSecOps Handbook...")
    doc = Document()
    
    # ตั้งค่าหน้ากระดาษ
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)  # A4
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # หน้าปก
    add_cover_page(doc)
    
    # สารบัญ
    add_toc_placeholder(doc)
    
    # ========== PART I: FOUNDATIONS ==========
    doc.add_heading("PART I — FOUNDATIONS & CORE FRAMEWORK", level=1)
    doc.add_paragraph()
    
    add_section(doc, "1. Advisory / คำเตือนและขอบเขต",
        "This document provides general guidelines for DevSecOps implementation. It is not a substitute for professional security audit or legal review. No warranty is provided.",
        "เอกสารนี้เป็นแนวทางทั่วไป ไม่ใช่คำแนะนำทางกฎหมายหรือการตรวจสอบความปลอดภัยแบบมืออาชีพ ผู้ใช้ต้องประเมินความเหมาะสมด้วยตนเอง ไม่มีการรับประกันผลลัพธ์ใดๆ ทั้งสิ้น")
    
    add_section(doc, "2. Scripts / สคริปต์และเครื่องมือ",
        "All scripts are provided under the license terms. Test in non-production environments first.",
        "สคริปต์ทั้งหมดจัดเตรียมไว้ตามเงื่อนไขใบอนุญาต ให้ทดสอบในสภาพแวดล้อมทดลองก่อนนำไปใช้งานจริง ดูตัวอย่างโค้ดใน Appendix A")
    
    add_section(doc, "3. License / สัญญาอนุญาต",
        "MIT / Apache-2.0 — See full license text in Appendix.",
        "เผยแพร่ภายใต้สัญญาอนุญาต MIT และ Apache-2.0 ดูรายละเอียดฉบับเต็มในภาคผนวก")
    
    add_section(doc, "4. Privacy Policy / นโยบายความเป็นส่วนตัว",
        "We process minimal necessary data. No personal data is stored without explicit consent.",
        "ประมวลผลข้อมูลเท่าที่จำเป็น ไม่จัดเก็บข้อมูลส่วนบุคคลโดยไม่ได้รับความยินยอมชัดเจน ปฏิบัติตาม GDPR, PDPA")
    
    add_section(doc, "5. Terms of Service / เงื่อนไขการใช้งาน",
        "Use at your own risk. Compliance with local laws is user responsibility.",
        "การใช้งานอยู่ในความรับผิดชอบของผู้ใช้ ต้องปฏิบัติตามกฎหมายที่บังคับใช้ในประเทศที่ตั้งระบบ")
    
    add_section(doc, "6. Policies / นโยบายความปลอดภัยและการเข้าถึง",
        "Role-based access control, least privilege, audit logging required.",
        "ใช้หลักการสิทธิน้อยที่สุด, แยกตามบทบาท, บันทึกการเข้าถึงทุกครั้ง เข้ารหัสข้อมูลสำคัญ")
    
    # ========== PART II: PLATFORM ==========
    doc.add_heading("PART II — PLATFORM & INTEGRATION COMPONENTS", level=1)
    doc.add_paragraph()
    
    add_section(doc, "7. Landing Page / หน้าแรกและภาพรวมระบบ",
        "Project overview, key features, architecture diagram, quick start guide.",
        "ภาพรวมโครงการ, คุณสมบัติหลัก, แผนผังสถาปัตยกรรม, คู่มือเริ่มต้นใช้งานอย่างรวดเร็ว")
    
    add_section(doc, "8. WebUI / ส่วนติดต่อผู้ใช้บนเว็บ",
        "Dashboard, settings, alerts, report generation.",
        "แดชบอร์ดภาพรวม, การตั้งค่าระบบ, การแจ้งเตือน, การสร้างรายงาน")
    
    add_section(doc, "9. Webhook / การตั้งค่าและการใช้งาน Webhook",
        "Endpoint configuration, payload format, authentication, signature verification.",
        "การตั้งค่า Endpoint, รูปแบบข้อมูลที่ส่ง, การตรวจสอบสิทธิ์ด้วยลายเซ็น")
    
    add_section(doc, "10. GitLab Flows / CI/CD Pipeline",
        "Pipeline architecture, runners, security scanning integration.",
        "โครงสร้าง Pipeline, Runners, การผสานการสแกนความปลอดภัย, เงื่อนไขการอนุมัติเผยแพร่")
    
    add_section(doc, "11. API & Spreadsheets / REST API และการเชื่อมต่อตาราง",
        "RESTful endpoints, authentication, rate limits, Google Sheets/Excel sync.",
        "รายการ Endpoint, การตรวจสอบสิทธิ์, ขีดจำกัดการเรียกใช้, การซิงค์ข้อมูลกับตาราง")
    
    add_section(doc, "12. SkillsBook / คู่มือทักษะและคำศัพท์",
        "Glossary of DevSecOps terms, capabilities, roles, best practices.",
        "คำศัพท์ทางเทคนิค, ความสามารถของระบบ, บทบาทหน้าที่, คู่มือแนวทางปฏิบัติที่ดี")
    
    # ========== PART III: MULTI-LANGUAGE ==========
    doc.add_heading("PART III — MULTI-LANGUAGE IMPLEMENTATION", level=1)
    doc.add_paragraph()
    
    languages = [
        ("13. Python", "SDK, automation scripts, API clients.", "ชุดเครื่องมือ, สคริปต์อัตโนมัติ, ตัวเชื่อมต่อ API"),
        ("14. Go (Golang)", "High-performance microservices, concurrent processing.", "บริการไมโครเซอร์วิสประสิทธิภาพสูง, การประมวลผลแบบขนาน"),
        ("15. JavaScript / TypeScript", "Frontend components, Node.js backend.", "ส่วนติดต่อผู้ใช้, ส่วนหลังด้วย Node.js"),
        ("16. Crystal", "Fast, safe, syntax-friendly.", "รวดเร็ว ปลอดภัย ไวยากรณ์อ่านง่าย"),
        ("17. Sigma", "Standardized detection rules, threat signatures.", "รูปแบบกฎการตรวจจับมาตรฐาน, ลายเซ็นภัยคุกคาม"),
        ("18. Rust", "Memory safety, zero-cost abstractions.", "ปลอดภัยหน่วยความจำ, ประสิทธิภาพสูงสุด"),
        ("19. Java", "Enterprise-grade, Spring Boot integration.", "มาตรฐานระดับองค์กร, ผสาน Spring Boot ได้สมบูรณ์"),
        ("20. C# / .NET", "Windows/Azure native, SDK for .NET ecosystem.", "ทำงานดั้งเดิมกับ Windows/Azure, ชุดเครื่องมือสำหรับระบบ .NET"),
    ]
    for title, en, th in languages:
        add_section(doc, title, en, th)
    
    # ========== PART IV: TOOLS & KNOWLEDGE ==========
    doc.add_heading("PART IV — TOOLS, STYLING & KNOWLEDGE SYSTEM", level=1)
    doc.add_paragraph()
    
    add_section(doc, "21. Tools & Styling / เครื่องมือและรูปแบบการเขียน",
        "Obsidian, Ollama, Claude, React/UI frameworks.",
        "ใช้ Obsidian จัดการเอกสาร, Ollama AI ในเครื่อง, ผสาน Claude, รูปแบบการเขียนด้วย React")
    
    add_section(doc, "22. Compliance Scripts / การตรวจสอบการปฏิบัติตามมาตรฐาน",
        "GDPR, ISO 27001, SOC2 automated validation.",
        "ตรวจสอบความสอดคล้องกับ GDPR, ISO 27001, SOC2 แบบอัตโนมัติ")
    
    add_section(doc, "23. Audit Trail Scripts / ระบบบันทึกและติดตามการเปลี่ยนแปลง",
        "Immutable logs, change tracking, export formats.",
        "บันทึกที่ไม่สามารถแก้ไขย้อนหลังได้, ติดตามการเปลี่ยนแปลงทั้งหมด")
    
    add_section(doc, "24. Automation Features / คุณสมบัติอัตโนมัติ",
        "DOCX/PDF export, Slack/Jira/email alerts, scheduled reports.",
        "ส่งออกไฟล์อัตโนมัติ, แจ้งเตือนผ่านช่องทางต่างๆ, สร้างรายงานตามกำหนดเวลา")
    
    add_section(doc, "25. Index / สารบัญและดัชนีคำสำคัญ",
        "Complete keyword index, table of contents.",
        "ดัชนีคำสำคัญฉบับสมบูรณ์, สารบัญ, คู่มืออ้างอิงด่วน")
    
    add_section(doc, "26. Debugging / การแก้ไขปัญหา",
        "Troubleshooting guide, log checklist, CI/CD diagnostics.",
        "คู่มือแก้ไขปัญหา, รายการตรวจสอบบันทึกระบบ, การวินิจฉัย Pipeline")
    
    add_section(doc, "27. Security Scan / การสแกนความปลอดภัย",
        "CodeQL, Snyk, Trivy, SAST/DAST, SBOM generation.",
        "สแกนโค้ด, ตรวจสอบความปลอดภัยแบบคงที่/ไดนามิก, สร้างรายการส่วนประกอบซอฟต์แวร์")
    
    add_section(doc, "28. Knowledge Section / ฐานความรู้",
        "CVSS scoring, threat modeling, incident response, risk assessment.",
        "ระบบให้คะแนนความรุนแรง, การสร้างแบบจำลองภัยคุกคาม, ขั้นตอนการจัดการเหตุการณ์")
    
    add_section(doc, "29. DataService / บริการข้อมูล",
        "Data API, CRUD operations, dashboard integration.",
        "API สำหรับเข้าถึงข้อมูล, การจัดการข้อมูล, เชื่อมต่อแดชบอร์ด")
    
    add_section(doc, "30. Databases / การออกแบบและจัดการฐานข้อมูล",
        "PostgreSQL, MongoDB, SQLite — schema design, backup, encryption.",
        "การออกแบบโครงสร้างข้อมูล, การสำรองข้อมูล, ระยะเวลาเก็บรักษา, การเข้ารหัส")
    
    add_section(doc, "31. AI-Driven Features / คุณสมบัติที่ขับเคลื่อนด้วย AI",
        "AI agents, auto-drafting, smart analysis, knowledge graph.",
        "ตัวแทนอัจฉริยะ, การร่างเอกสารอัตโนมัติ, การวิเคราะห์อัจฉริยะ, กราฟความรู้")
        # ========== PART IV EXTENDED: COLLABORATION, AI & KNOWLEDGE ==========
    doc.add_heading("PART IV — EXTENDED: Collaboration, AI & Knowledge Systems", level=1)
    doc.add_paragraph()

    add_section(doc, "32. Contribution / คู่มือการมีส่วนร่วม",
        "Guidelines for contributors, code standards, pull request process, code review workflow, CLA/DCO requirements.",
        "แนวทางผู้มีส่วนร่วม, มาตรฐานโค้ด, ขั้นตอน Pull Request, การตรวจสอบโค้ด, ข้อกำหนด CLA/DCO — ดูรายการตรวจสอบใน Appendix D")

    add_section(doc, "33. Collaboration / การทำงานร่วมกันและการจัดการทีม",
        "Team workflows, role definitions, communication channels, review cycles, cross-team coordination, access governance.",
        "ขั้นตอนการทำงานร่วมกัน, บทบาทหน้าที่, ช่องทางสื่อสาร, รอบการตรวจสอบ, การประสานงานข้ามทีม, การจัดการสิทธิ์")

    add_section(doc, "34. Feature Request / กระบวนการเสนอคุณสมบัติใหม่",
        "Proposal workflow, RFC template, prioritization framework, feedback loop, version planning, deprecation policy.",
        "ขั้นตอนเสนอฟีเจอร์, แบบฟอร์ม RFC, กรอบจัดลำดับความสำคัญ, วงรอบรับฟังความคิดเห็น, การวางแผนเวอร์ชัน, นโยบายเลิกใช้งาน")

    add_section(doc, "35. GitHub Copilot Run Sheets / คู่มือการใช้ AI คู่เขียนโค้ด",
        "Prompts & workflows for GitHub Copilot, code generation patterns, test automation, refactoring, security-aware prompts, cost & token control.",
        "ชุดคำสั่งและขั้นตอนการใช้งาน GitHub Copilot, รูปแบบการสร้างโค้ด, การสร้างทดสอบ, การปรับปรุงโค้ด, คำสั่งเช็คความปลอดภัย, การควบคุมค่าและโทเคน")

    add_section(doc, "36. Knowledge Platforms / ระบบฐานความรู้และเอกสาร",
        "Obsidian, MkDocs, GitBook, Notion integration — knowledge graph, bidirectional links, search, version sync, offline access, multi-platform publishing.",
        "การผสาน Obsidian, MkDocs, GitBook, Notion — กราฟความรู้, ลิงก์สองทาง, การค้นหา, การซิงค์เวอร์ชัน, การเข้าถึงแบบออฟไลน์, การเผยแพร่หลายแพลตฟอร์ม")


    # ========== APPENDICES ==========
    doc.add_page_break()
    doc.add_heading("APPENDIX — ภาคผนวก", level=1)
    doc.add_paragraph()
    
    appendices = [
        ("Appendix A — Full Scripts Collection", "รวบรวมสคริปต์ทั้งหมดในทุกภาษา — Python, Go, JavaScript, Rust, Java, C#, Crystal"),
        ("Appendix B — YAML Configurations", "ไฟล์กำหนดการทำงาน, Pipeline, นโยบายทั้งหมดในรูปแบบ YAML"),
        ("Appendix C — API Reference", "เอกสารอ้างอิง API ฉบับเต็ม, ตัวอย่างคำขอและคำตอบ"),
        ("Appendix D — Compliance Checklists", "รายการตรวจสอบความสอดคล้องกับ GDPR, ISO 27001, SOC2"),
        ("Appendix E — Error Codes & Troubleshooting", "ตารางรหัสข้อผิดพลาด, ปัญหาที่พบบ่อย, วิธีแก้ไข"),
        ("Appendix F — Glossary", "คำศัพท์ทางเทคนิค คำอธิบาย ทั้งภาษาไทยและอังกฤษ"),
    ]
    for title, desc in appendices:
        add_section(doc, title, desc, "", level=2)
    
    # ========== INDEX ==========
    add_index(doc)
    
    # บันทึกไฟล์
    doc.save(OUTPUT_FILE)
    print(f"✅ สร้างไฟล์สำเร็จ: {OUTPUT_FILE}")
    print("📖 เปิดไฟล์ใน Word → คลิกขวาที่สารบัญ → Update Field เพื่ออัพเดทเลขหน้า")

if __name__ == "__main__":
    main()
